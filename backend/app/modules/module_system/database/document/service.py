"""
文档服务层
"""
import os
import tempfile
from typing import Dict, Any

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from backend.app.common.core.exceptions import CustomException
from backend.app.common.core.logger import log
from backend.app.modules.module_system.auth.schema import AuthSchema
from backend.app.modules.module_system.database.knowledge.crud import KnowledgeCRUD
from .crud import KnowledgeFileCRUD, KnowledgeVectorCRUD
from .schema import (
    DocumentCreateSchema,
    DocumentOutSchema,
    DocumentUpdateSchema, MetadataSchema
)


class DocumentService:
    """文档服务"""

    @classmethod
    async def create_service(
        cls,
        auth: AuthSchema,
        knowledge_uuid: str,
        data: DocumentCreateSchema
    ) -> Dict[str, Any]:
        """添加文档切片到知识库
        三表两库架构全流程：
        1. 校验：知识库存在 + content 非空
        2. 保存：元数据到 document 关系表（status=0 处理中）
        3. 向量化：计算 content 向量并插入 kb_{id} 向量表
        4. 更新：document 表的 vector_id + status=1（成功）
        5. 统计：更新 knowledge 表的 document_count
        """
        kb_crud = KnowledgeCRUD(auth=auth)
        file_crud = KnowledgeFileCRUD(auth=auth)

        # 1. 校验：知识库必须存在
        kb = await kb_crud.get(uuid=knowledge_uuid)
        if not kb:
            raise CustomException(msg="知识库不存在")

        # 2. 校验：content 不能为空
        if not data.content or len(data.content.strip()) == 0:
            raise CustomException(msg="文档内容不能为空")

        # 3. 保存文件元数据到 PostgreSQL 关系表（不存 content）
        #    初始状态：status=0 处理中
        doc_data = data.model_dump(exclude={"content", "knowledge_uuid"})
        doc_data["knowledge_id"] = kb.id
        doc_data["status"] = 0  # 处理中
        doc = await file_crud.create(doc_data)

        vector_id = None
        try:
            # 4. 添加到 PGVector 向量库（自动计算嵌入向量）
            vector_crud = KnowledgeVectorCRUD(kb_id=kb.id)

            # 使用 MetadataSchema 规范化元数据（单切片文档 page=1, chunk_index=0）

            metadata = MetadataSchema(
                knowledge_id=kb.id,
                document_id=doc.id,
                file_name=data.file_name or data.title,
                page=1,  # 单切片文档页码为1
                chunk_index=0,  # 单切片文档序号为0
            )

            # 创建 LangChain Document 对象
            document = Document(
                page_content=data.content,
                metadata={
                    **metadata.model_dump(),  # 标准字段
                    "title": data.title,
                    "source": data.source,
                    **(data.meta_data or {})  # 扩展元数据
                }
            )

            # 插入向量到 PGVector，返回向量ID
            ids = vector_crud.add_documents([document])
            vector_id = ids[0] if ids else None

            if not vector_id:
                raise CustomException(msg="向量插入失败")

            # 5. 更新文档记录：保存 vector_id + 状态改为成功
            await file_crud.update(doc.id, {
                "vector_id": vector_id,
                "status": 1  # 成功
            })

            # 重新读取完整文档信息
            updated_doc = await file_crud.get(id=doc.id)
            log.info(f"文档添加成功: document_id={doc.id}, vector_id={vector_id}")

        except Exception as e:
            # 向量插入失败：标记文档为失败状态
            await file_crud.update(doc.id, {"status": 2})  # 失败
            log.warning(f"文档添加失败: document_id={doc.id}, error={str(e)}")
            raise CustomException(msg=f"文档向量化失败: {str(e)}")

        # 使用 DocumentOutSchema 标准化返回格式
        from sqlalchemy import inspect as sa_inspect
        doc_data = {c.key: getattr(updated_doc, c.key) for c in sa_inspect(updated_doc).mapper.column_attrs}
        doc_data["knowledge_uuid"] = knowledge_uuid
        return DocumentOutSchema.model_validate(doc_data).model_dump(mode='json')

    @classmethod
    async def delete_service(
        cls,
        auth: AuthSchema,
        knowledge_uuid: str,
        doc_id: int
    ) -> Dict[str, Any]:
        """删除文档
        - 软删除关系表
        - 删除向量表中的对应向量
        """
        kb_crud = KnowledgeCRUD(auth=auth)
        file_crud = KnowledgeFileCRUD(auth=auth)

        # 1. 校验知识库存在
        kb = await kb_crud.get(uuid=knowledge_uuid)
        if not kb:
            raise CustomException(msg="知识库不存在")

        # 2. 校验文档存在
        doc = await file_crud.get(id=doc_id)
        if not doc:
            raise CustomException(msg="文档不存在")

        # 3. 软删除文档
        await file_crud.delete([doc_id])

        # 4. 删除向量表中的对应向量
        vector_crud = KnowledgeVectorCRUD(kb_id=kb.id)
        deleted_vector_count = vector_crud.delete_by_file_id(doc_id)
        log.info(f"删除文档向量: doc_id={doc_id}, vector_count={deleted_vector_count}")

        return {
            "deleted_id": doc_id,
            "kb_id": kb.id,
            "deleted_vector_count": deleted_vector_count
        }

    @classmethod
    async def update_service(
        cls,
        auth: AuthSchema,
        doc_id: int,
        data: DocumentUpdateSchema
    ) -> Dict[str, Any]:
        """更新文档元信息
        - 仅更新元信息，不修改向量内容
        - 支持迁移知识库（修改 knowledge_id）
        """
        # 1. 校验文档存在
        doc = await KnowledgeFileCRUD(auth=auth).get(id=doc_id)
        if not doc:
            raise CustomException(msg="文档不存在")

        # 2. 获取当前知识库（用于返回 knowledge_uuid）
        current_kb = await KnowledgeCRUD(auth=auth).get(id=doc.knowledge_id)
        if not current_kb:
            raise CustomException(msg="当前知识库不存在")

        # 3. 准备更新数据（使用 exclude_unset=True 支持部分更新）
        update_data = data.model_dump(exclude={"document_id", "knowledge_uuid"}, exclude_unset=True)

        # 4. 如果指定了新的 knowledge_uuid，需要迁移知识库
        if data.knowledge_uuid and data.knowledge_uuid != current_kb.uuid:
            target_kb = await KnowledgeCRUD(auth=auth).get(uuid=data.knowledge_uuid)
            if not target_kb:
                raise CustomException(msg="目标知识库不存在")
            update_data["knowledge_id"] = target_kb.id
            target_kb_uuid = target_kb.uuid
        else:
            target_kb_uuid = current_kb.uuid

        # 5. 更新文档
        updated_doc = await KnowledgeFileCRUD(auth=auth).update(doc_id, update_data)

        # 6. 转换为字典并添加 knowledge_uuid
        from sqlalchemy import inspect as sa_inspect
        doc_data = {c.key: getattr(updated_doc, c.key) for c in sa_inspect(updated_doc).mapper.column_attrs}
        doc_data["knowledge_uuid"] = target_kb_uuid

        return DocumentOutSchema.model_validate(doc_data).model_dump(mode='json')


    @classmethod
    async def detail_service(
        cls,
        auth: AuthSchema,
        knowledge_uuid: str,
        doc_id: int
    ) -> Dict[str, Any]:
        """获取文档详情"""
        # 校验知识库存在
        kb = await KnowledgeCRUD(auth=auth).get(uuid=knowledge_uuid)
        if not kb:
            raise CustomException(msg="知识库不存在")

        doc = await KnowledgeFileCRUD(auth=auth).get(id=doc_id)
        if not doc:
            raise CustomException(msg="文档不存在")

        # 转换为字典并添加 knowledge_uuid
        from sqlalchemy import inspect as sa_inspect
        doc_data = {c.key: getattr(doc, c.key) for c in sa_inspect(doc).mapper.column_attrs}
        doc_data["knowledge_uuid"] = kb.uuid

        return DocumentOutSchema.model_validate(doc_data).model_dump(mode='json')

    @classmethod
    async def upload_file_service(
        cls,
        auth: AuthSchema,
        knowledge_uuid: str,
        file_content: bytes,
        file_name: str,
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ) -> Dict[str, Any]:
        """
        上传文件并自动切片存入知识库

        参数:
            knowledge_uuid: 知识库UUID
            file_content: 文件内容字节
            file_name: 文件名
            chunk_size: 切片大小
            chunk_overlap: 切片重叠

        返回:
            上传结果
        """
        kb_crud = KnowledgeCRUD(auth=auth)
        file_crud = KnowledgeFileCRUD(auth=auth)

        # 1. 获取知识库
        kb = await kb_crud.get(uuid=knowledge_uuid)
        if not kb:
            raise CustomException(msg="知识库不存在")

        # 2. 保存临时文件
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file_name)[1]) as tmp:
            tmp.write(file_content)
            tmp_path = tmp.name

        try:
            # 3. 加载文件内容
            content = DocumentService.load_file_content(tmp_path, file_name)
            if not content or len(content.strip()) == 0:
                raise CustomException(msg="文件内容为空")

            # 4. 文本切片
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
                length_function=len
            )
            chunks = text_splitter.split_text(content)

            if not chunks:
                raise CustomException(msg="切片后没有生成有效切片")

            # 5. 保存文件元数据（主文件记录，初始状态：处理中）
            file_size = len(file_content)
            total_chunks = len(chunks)
            doc_data = {
                "title": file_name,
                "file_name": file_name,
                "file_size": file_size,
                "chunk_count": total_chunks,
                "source": file_name,
                "status": 0,  # 处理中
                "knowledge_id": kb.id,
            }
            doc = await file_crud.create(doc_data)

            try:
                # 6. 使用 MetadataSchema 规范化元数据，创建 Document 对象列表
                from .schema import MetadataSchema
                documents = []
                for i, chunk in enumerate(chunks):
                    metadata = MetadataSchema(
                        knowledge_id=kb.id,
                        document_id=doc.id,
                        file_name=file_name,
                        page=(i // 50) + 1,  # 每50切片算1页
                        chunk_index=i,
                    )
                    documents.append(Document(
                        page_content=chunk,
                        metadata={
                            **metadata.model_dump(),  # 标准字段
                            "title": file_name,
                            "source": file_name,
                            "total_chunks": total_chunks,  # 扩展字段
                        }
                    ))

                # 7. 批量插入向量到 PGVector
                vector_crud = KnowledgeVectorCRUD(kb_id=kb.id)
                vector_ids = vector_crud.add_documents(documents)

                if not vector_ids or len(vector_ids) == 0:
                    raise CustomException(msg="向量批量插入失败")

                # 8. 更新文件记录：保存第一个 vector_id + 状态改为成功
                await file_crud.update(doc.id, {
                    "vector_id": vector_ids[0],
                    "status": 1  # 成功
                })

                # 重新读取完整文档信息
                updated_doc = await file_crud.get(id=doc.id)
                log.info(f"文件上传成功: {file_name}, 切片数: {total_chunks}, 向量数: {len(vector_ids)}")

                # 使用 DocumentOutSchema 标准化返回格式
                from sqlalchemy import inspect as sa_inspect
                doc_data = {c.key: getattr(updated_doc, c.key) for c in sa_inspect(updated_doc).mapper.column_attrs}
                doc_data["knowledge_uuid"] = knowledge_uuid
                return DocumentOutSchema.model_validate(doc_data).model_dump(mode='json')

            except Exception as e:
                # 向量插入失败：标记文档为失败状态
                await file_crud.update(doc.id, {"status": 2})  # 失败
                log.warning(f"文件上传失败: {file_name}, document_id={doc.id}, error={str(e)}")
                raise CustomException(msg=f"文件向量化失败: {str(e)}")

        finally:
            # 清理临时文件
            try:
                os.unlink(tmp_path)
            except:
                pass

    @staticmethod
    def load_file_content(file_path: str, file_name: str) -> str:
        """根据文件扩展名加载文件内容"""
        ext = os.path.splitext(file_name)[1].lower()

        if ext in ['.txt', '.md', '.py', '.js', '.json', '.yaml', '.yml', '.xml', '.html', '.css', '.java', '.go',
                   '.cpp', '.c', '.h', '.rs']:
            return DocumentService.load_text_file(file_path)
        elif ext == '.pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() or ""
                return text
            except ImportError:
                raise CustomException(msg="PDF 支持需要安装 PyPDF2：pip install PyPDF2")
        elif ext in ['.docx', '.doc']:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for para in doc.paragraphs:
                    text += para.text + "\n"
                return text
            except ImportError:
                raise CustomException(msg="Word 文档支持需要安装 python-docx：pip install python-docx")
        else:
            # 默认按文本加载
            return DocumentService.load_text_file(file_path)

    @staticmethod
    # 文档加载器
    def load_text_file(file_path: str) -> str:
        """加载纯文本文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
            except UnicodeDecodeError:
                continue
        raise CustomException(msg="文件编码不支持，仅支持 UTF-8、GBK、GB2312 编码")



